import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import dns.resolver
from typing import Optional
import re
import pandas as pd
import docx
import os

class AutoMailer:
    """Class to send emails using SMTP."""
    def __init__(self, username: str, password: str, smtp_server: Optional[str]=None, smtp_port: Optional[int] = 587):
        if self._check_mail(username) is False:
            raise ValueError("Invalid email address format.")

        if smtp_server is None:
            self.smtp_server = self._infer_smtp_server(username)
            print(f"Inferred SMTP server: {self.smtp_server}")
        else:
            self.smtp_server = smtp_server

        if 0 >= smtp_port or smtp_port > 65_535:
            raise ValueError("SMTP port must be between 1 and 65535.")
        
        # Default SMTP port is 587 for TLS
        if smtp_port is None:
            self.smtp_port = 587
        else:
            self.smtp_port = smtp_port
            
        self.username = username
        self.password = password
    
    @staticmethod
    def _check_mail(username: str) -> bool:
        """
        Function to check if the email address is vaild.
        Args:
            username (str): The email address to check.
        Returns:
            bool: True if the email address is valid, False otherwise.
        """
        email_regex = r'^[a-zA-z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return re.match(email_regex, username) is not None
    
    @staticmethod
    def _infer_smtp_server(username: str) -> str:
        """
        Function to infer the SMTP server from the email address.
        Args:
            username (str): The email address.
        Returns:
            str: The inferred SMTP server.
        """
        domain = username.split('@')[-1].strip()
        print(f"Inferring SMTP server for domain: {domain}")
        try:
            answers = dns.resolver.resolve(domain, 'MX')
            mx_records = []

            for rdata in answers:
                mx_host = str(rdata.exchange).rstrip('.')
                mx_records.append((rdata.preference, mx_host))
            
            mx_records.sort()
            priority, mx_host = mx_records[0]

            mx_parts = mx_host.split('.')
            if mx_parts[0].startswith('mx'):
                
                mx_parts[0] = 'smtp'
            smtp_server = '.'.join(mx_parts)

            return smtp_server
        except Exception as e:
            raise RuntimeError(f"Could not resolve MX records for domain {domain}: {e}")
    
    def _render_template(self, row_data: dict) -> str:
        """
        Render the docx template by replacing placeholders {{ column_name }} with actual data.
        Returns final plain text message.
        """
        doc = docx.Document(self.template_path)
        placeholder_pattern = re.compile(r"\{\{\s*(\w+)\s*\}\}")

        for paragraph in doc.paragraphs:
            full_text = "".join(run.text for run in paragraph.runs)
            matches = placeholder_pattern.findall(full_text)
            
            if matches:
                for match in matches:
                    if match in row_data:
                        full_text = re.sub(r"\{\{\s*" + re.escape(match) + r"\s*\}\}", str(row_data[match]), full_text)

                # Clear all runs in paragraph
                for run in paragraph.runs:
                    run.text = ""

                # Add a single new run with replaced text (you can also style it if needed)
                paragraph.runs[0].text = full_text

        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        
    def send_mail(self, to_address: str, subject: str, body: str) -> bool:
        """
            Function to send an email.
            Args:
                to_address (str): The recipient's email address.
                subject (str): The subject of the email.
                body (str): The body of the email.
            Returns:
                bool: True if the email was sent successfully, False otherwise.
        """
        msg = MIMEMultipart()
        msg['From'] = self.username
        msg['To'] = to_address
        msg['Subject'] = subject
        
        msg.attach(MIMEText(body, 'plain'))

        try:
            if self.smtp_port == 465:
                #SSL
                with smtplib.SMTP_SSL(self.smtp_server, self.smtp_port) as server:
                    server.login(self.username, self.password)
                    server.sendmail(self.username, to_address, msg.as_string())
            else:
                #TLS
                with smtplib.SMTP(self.smtp_server, self.smtp_port) as server:
                    server.starttls()
                    server.login(self.username, self.password)
                    server.sendmail(self.username, to_address, msg.as_string())

                print(f"Email sucessfully sent to {to_address}")
                return True
            
        except Exception as e:
            print(f"Failed to send email: {e}")
            return False

class BulkMailer(AutoMailer):
    """Class to send bulk emails."""
    def __init__(self, username: str, password: str, csv_path: str, smtp_server: Optional[str]=None, smtp_port: Optional[int] = 587, template_path: Optional[str]=None):
        super().__init__(username, password, smtp_server, smtp_port)
        self.csv_path = csv_path
        self.template_path = os.path.abspath(template_path) if template_path else None
    
    def bulk_mail(self, 
              mail_to_column: str, 
              subject_column: Optional[str]=None, 
              body_column: Optional[str]=None,
              subject: Optional[str]=None):
        """
        Send bulk emails using columns or docx template.
        Args:
            mail_to_column (str): Column with recipient emails.
            subject_column (str): Column with subject (if not using template).
            body_column (str): Column with body (if not using template).
            subject (str): Subject to use when using template (if no subject column).
        """
        if self.csv_path.endswith('.csv'):
            df = pd.read_csv(self.csv_path)
        elif self.csv_path.endswith('.xlsx'):
            df = pd.read_excel(self.csv_path)
        else:
            raise ValueError("File format not supported. Please use CSV or Excel files.")

        for _, row in df.iterrows():
            to_address = row[mail_to_column]

            # When using template
            if self.template_path:
                # Render email body from template using row
                body = self._render_template(row)

                # Subject: either from column or passed subject argument
                if subject_column and subject_column in row:
                    email_subject = row[subject_column]
                else:
                    email_subject = subject if subject else f"Notification for {row.get('Name', to_address)}"

                self.send_mail(to_address, email_subject, body)

            # When NOT using template
            else:
                email_subject = row[subject_column] if subject_column else subject
                body = row[body_column] if body_column else ""
                self.send_mail(to_address, email_subject, body)
